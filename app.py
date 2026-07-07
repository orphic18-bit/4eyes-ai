"""
4eyes.ai — Streamlit Frontend
Contract Intelligence Platform
"""

import streamlit as st
import tempfile
import os
import sys
import io
import html
from pathlib import Path

# ── Page config (must be first Streamlit call) ────────────────────────────────
st.set_page_config(
    page_title="4eyes.ai — Contract Intelligence",
    page_icon="👁️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Styling ───────────────────────────────────────────────────────────────────
st.markdown("""
<style>
/* ── Fonts ── */
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Mono:wght@400;500&family=DM+Sans:wght@300;400;500&display=swap');

/* ── Base ── */
html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
    background-color: #0A0A0F;
    color: #E8E4DC;
}

/* Kill default Streamlit chrome */
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 2rem 3rem 4rem; max-width: 1100px; }

/* ── Brand header ── */
.brand-header {
    display: flex;
    align-items: baseline;
    gap: 0.5rem;
    margin-bottom: 0.25rem;
}
.brand-name {
    font-family: 'DM Serif Display', serif;
    font-size: 2.6rem;
    color: #E8E4DC;
    letter-spacing: -0.02em;
    line-height: 1;
}
.brand-dot {
    font-family: 'DM Serif Display', serif;
    font-size: 2.6rem;
    color: #C8F55A;
    line-height: 1;
}
.brand-tagline {
    font-family: 'DM Mono', monospace;
    font-size: 0.72rem;
    color: #6B6860;
    letter-spacing: 0.12em;
    text-transform: uppercase;
    margin-bottom: 2.5rem;
}

/* ── Divider ── */
.hairline {
    border: none;
    border-top: 1px solid #1E1E26;
    margin: 2rem 0;
}

/* ── Section labels ── */
.section-label {
    font-family: 'DM Mono', monospace;
    font-size: 0.68rem;
    letter-spacing: 0.14em;
    text-transform: uppercase;
    color: #6B6860;
    margin-bottom: 0.75rem;
}

/* ── Input mode pills ── */
.stRadio > div {
    display: flex;
    gap: 0.5rem;
    flex-direction: row !important;
}
.stRadio > div > label {
    background: #13131A;
    border: 1px solid #2A2A35;
    border-radius: 4px;
    padding: 0.4rem 1rem;
    font-size: 0.82rem;
    cursor: pointer;
    transition: all 0.15s;
    color: #9A9690 !important;
}
.stRadio > div > label:has(input:checked) {
    background: #1A1A24;
    border-color: #C8F55A;
    color: #C8F55A !important;
}

/* ── File uploader ── */
[data-testid="stFileUploader"] {
    background: #0E0E16;
    border: 1px dashed #2A2A35;
    border-radius: 8px;
    padding: 1.5rem;
    transition: border-color 0.2s;
}
[data-testid="stFileUploader"]:hover {
    border-color: #C8F55A40;
}

/* ── Text area ── */
.stTextArea > div > div > textarea {
    background: #0E0E16 !important;
    border: 1px solid #2A2A35 !important;
    border-radius: 8px;
    color: #E8E4DC !important;
    font-family: 'DM Mono', monospace !important;
    font-size: 0.82rem;
    line-height: 1.6;
    padding: 1rem;
    min-height: 220px;
}
.stTextArea > div > div > textarea:focus {
    border-color: #C8F55A60 !important;
    box-shadow: 0 0 0 1px #C8F55A30 !important;
}

/* ── Analyse button ── */
.stButton > button {
    background: #C8F55A;
    color: #0A0A0F;
    font-family: 'DM Mono', monospace;
    font-size: 0.82rem;
    font-weight: 500;
    letter-spacing: 0.06em;
    text-transform: uppercase;
    border: none;
    border-radius: 4px;
    padding: 0.65rem 2.2rem;
    cursor: pointer;
    transition: background 0.15s, transform 0.1s;
    width: auto;
}
.stButton > button:hover {
    background: #D8FF6A;
    transform: translateY(-1px);
}
.stButton > button:active {
    transform: translateY(0);
}
.stButton > button:disabled {
    background: #2A2A35;
    color: #4A4A55;
    cursor: not-allowed;
    transform: none;
}

/* ── Spinner ── */
.stSpinner > div {
    border-color: #C8F55A transparent transparent transparent;
}

/* ── Report output ── */
.report-wrapper {
    background: #0E0E16;
    border: 1px solid #1E1E26;
    border-radius: 8px;
    padding: 2rem 2.5rem;
    margin-top: 1.5rem;
}

/* Section headers inside report */
.report-wrapper h2 {
    font-family: 'DM Serif Display', serif;
    font-size: 1.3rem;
    color: #E8E4DC;
    margin: 1.8rem 0 0.6rem;
    padding-bottom: 0.4rem;
    border-bottom: 1px solid #1E1E26;
}
.report-wrapper h3 {
    font-family: 'DM Sans', sans-serif;
    font-size: 0.9rem;
    font-weight: 500;
    color: #C8F55A;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    margin: 1.2rem 0 0.4rem;
}
.report-wrapper p, .report-wrapper li {
    font-size: 0.88rem;
    color: #B0ACA4;
    line-height: 1.75;
}
.report-wrapper strong {
    color: #E8E4DC;
    font-weight: 500;
}
.report-wrapper code {
    font-family: 'DM Mono', monospace;
    font-size: 0.78rem;
    background: #13131A;
    padding: 0.1em 0.4em;
    border-radius: 3px;
    color: #C8F55A;
}
.report-wrapper blockquote {
    border-left: 2px solid #C8F55A40;
    margin: 0.8rem 0;
    padding: 0.4rem 1rem;
    color: #7A7870;
    font-style: italic;
}

/* Risk badges */
.risk-high   { color: #FF6B6B; font-weight: 600; }
.risk-medium { color: #FFB347; font-weight: 600; }
.risk-low    { color: #C8F55A; font-weight: 600; }

/* ── Meta bar ── */
.meta-bar {
    display: flex;
    gap: 2rem;
    padding: 0.75rem 0;
    border-bottom: 1px solid #1E1E26;
    margin-bottom: 1.5rem;
}
.meta-item {
    font-family: 'DM Mono', monospace;
    font-size: 0.7rem;
    color: #6B6860;
    letter-spacing: 0.06em;
    text-transform: uppercase;
}
.meta-value {
    color: #9A9690;
    font-weight: 500;
}

/* ── Error / warning boxes ── */
.stAlert {
    background: #13131A;
    border-radius: 6px;
    border: 1px solid #2A2A35;
}

/* ── Download button ── */
.stDownloadButton > button {
    background: transparent;
    color: #6B6860;
    font-family: 'DM Mono', monospace;
    font-size: 0.72rem;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    border: 1px solid #2A2A35;
    border-radius: 4px;
    padding: 0.4rem 1.2rem;
    transition: all 0.15s;
}
.stDownloadButton > button:hover {
    border-color: #4A4A55;
    color: #E8E4DC;
}

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 4px; }
::-webkit-scrollbar-track { background: #0A0A0F; }
::-webkit-scrollbar-thumb { background: #2A2A35; border-radius: 2px; }
</style>
""", unsafe_allow_html=True)


# ── Brand header ──────────────────────────────────────────────────────────────
st.markdown("""
<div class="brand-header">
    <span class="brand-name">4eyes</span><span class="brand-dot">.</span><span class="brand-name">ai</span>
</div>
<div class="brand-tagline">Plain-English Contract Audits &nbsp;·&nbsp; Built for Freelancers &amp; Solo Founders</div>
""", unsafe_allow_html=True)

st.markdown('<hr class="hairline">', unsafe_allow_html=True)


# ── Helpers ───────────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError as err:
        raise RuntimeError(
            "pdfplumber is not installed. Run: pip install pdfplumber"
        ) from err

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n\n".join(pages).strip()
    except Exception as err:
        raise RuntimeError(
            f"PDF text extraction failed. Please paste the contract text manually instead. Details: {err}"
        ) from err


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX using python-docx."""
    try:
        import docx
    except ImportError as err:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from err

    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        # Some contracts put commercial terms or signature blocks inside tables.
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n\n".join(parts).strip()
    except Exception as err:
        raise RuntimeError(
            f"DOCX text extraction failed. Please paste the contract text manually instead. Details: {err}"
        ) from err


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from a plain TXT file."""
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue

    raise RuntimeError(
        "TXT text extraction failed. Please paste the contract text manually instead."
    )


def extract_text_from_uploaded_file(uploaded_file) -> str:
    """Route an uploaded file to the correct text extractor."""
    file_bytes = uploaded_file.getvalue()
    ext = Path(uploaded_file.name).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    if ext == ".docx":
        return extract_text_from_docx(file_bytes)
    if ext == ".txt":
        return extract_text_from_txt(file_bytes)

    raise RuntimeError(
        "Unsupported file type. Please upload a PDF, DOCX, or TXT file, or paste the contract text manually."
    )


def run_analysis(contract_text: str) -> str:
    """
    Import and run the contract_analyzer pipeline.

    contract_analyzer.py must expose one of:
      - analyze(text: str) -> str          (returns formatted report string)
      - analyze(text: str) -> dict         (returns structured dict)
      - ContractAnalyzer().analyze(text)   (class-based)

    Adjust the call below to match your actual API.
    """
    # Ensure the project root is on sys.path so contract_analyzer can be imported
    project_root = Path(__file__).parent
    if str(project_root) not in sys.path:
        sys.path.insert(0, str(project_root))

    try:
        import contract_analyzer  # noqa: PLC0415
    except ModuleNotFoundError:
        raise RuntimeError(
            "contract_analyzer.py not found. "
            "Make sure it lives in the same directory as app.py."
        )

    # ── Try the most common API shapes ────────────────────────────────────────
    if hasattr(contract_analyzer, "analyze"):
        result = contract_analyzer.analyze(contract_text)
    elif hasattr(contract_analyzer, "run"):
        result = contract_analyzer.run(contract_text)
    elif hasattr(contract_analyzer, "ContractAnalyzer"):
        analyzer = contract_analyzer.ContractAnalyzer()
        result = analyzer.analyze(contract_text)
    elif hasattr(contract_analyzer, "main"):
        # Some scripts write to stdout; capture it
        old_stdout = sys.stdout
        sys.stdout = buffer = io.StringIO()
        try:
            contract_analyzer.main(contract_text)
        finally:
            sys.stdout = old_stdout
        result = buffer.getvalue()
    else:
        raise RuntimeError(
            "Could not find a recognised entry-point in contract_analyzer.py. "
            "Expected: analyze(), run(), ContractAnalyzer().analyze(), or main()."
        )

    # Normalise to string
    if isinstance(result, dict):
        import json
        return json.dumps(result, indent=2)
    return str(result)


# ── Sample contract for zero-friction testing ────────────────────────────────
SAMPLE_CONTRACT = """INDEPENDENT CONTRACTOR SERVICES AGREEMENT

This Independent Contractor Services Agreement ("Agreement") is entered into as of June 15, 2026, by and between Nexus Global Cloud Solutions, Inc., a Delaware corporation ("Company" or "Nexus"), and The Undersigned ("Contractor").

1. Services and Deliverables. Contractor agrees to perform the software development and consulting services ("Services") and provide the deliverables ("Deliverables") as mutually agreed upon in writing by the parties. Contractor shall perform the Services in a professional and workmanlike manner, in accordance with highest industry standards.

2. Compensation and Payment Terms. Company shall pay Contractor the agreed-upon fees for the Services. Contractor shall submit invoices monthly. Company agrees to pay all undisputed invoices within Net-60 days of receipt, subject to Company's Final Acceptance of the Deliverables. "Final Acceptance" shall mean the Deliverables have been deployed to Company's production environment and Company has confirmed, in its sole subjective discretion, that no further revisions, bug fixes, or modifications are required. Company reserves the right of offset, and may withhold payment of any invoice to offset any damages, claims, or suspected breaches of this Agreement or any other agreement between Company and Contractor.

3. Intellectual Property and Pre-Existing Rights. The parties agree that Contractor is an independent entity. Contractor shall retain all right, title, and interest in and to any pre-existing code, libraries, frameworks, or tools that Contractor developed prior to this Agreement ("Pre-Existing IP"). All new work product created specifically for the Company under this Agreement shall be considered a "work made for hire" and shall be the sole property of the Company.

4. Term and Termination. This Agreement shall commence on the date first written above and continue until the Services are completed. Company may terminate this Agreement at any time, for any reason or no reason, immediately upon written notice to Contractor. Contractor may only terminate this Agreement in the event of a material breach by Company, provided that Contractor gives Company thirty (30) days prior written notice specifying the breach, and Company fails to cure such breach within the thirty (30) day period.

5. Representations and Warranties. Contractor represents and warrants that the Services and Deliverables will not infringe upon the intellectual property rights of any third party. Contractor agrees to defend, indemnify, and hold Company harmless from any claims, damages, or expenses arising from a breach of this warranty.

6. Audit Rights. During the term of this Agreement and for a period of two (2) years thereafter, Company or its designated representatives shall have the right, upon three (3) business days' notice, to inspect and audit Contractor's books, records, computer systems, and personal devices used in connection with the Services, to ensure compliance with the terms of this Agreement and the security of Company data.

7. Miscellaneous and General Provisions. This Agreement constitutes the entire understanding between the parties. If any provision is found to be unenforceable, the remaining provisions shall survive. Notwithstanding anything to the contrary in Section 3 or elsewhere in this Agreement, if Contractor incorporates any Pre-Existing IP into the Deliverables, Contractor hereby assigns and transfers all ownership rights, titles, and interests in that specific Pre-Existing IP to the Company globally and in perpetuity.

IN WITNESS WHEREOF, the parties have executed this Agreement.
Nexus Global Cloud Solutions, Inc. By: ______________________ Title: Director of Vendor Relations
Contractor: By: ______________________ Title: Independent Contractor"""


def _load_sample_contract() -> None:
    st.session_state["contract_text_input"] = SAMPLE_CONTRACT


# ── Input section ─────────────────────────────────────────────────────────────
st.markdown('<div class="section-label">UPLOAD CONTRACT (PDF, DOCX, or TXT)</div>', unsafe_allow_html=True)

uploaded = st.file_uploader(
    label="UPLOAD CONTRACT (PDF, DOCX, or TXT)",
    type=["pdf", "docx", "txt"],
    label_visibility="collapsed",
    help="PDF, DOCX, or TXT, up to 50 MB",
)

# Privacy reassurance — directly below the upload zone
st.markdown(
    '<div style="font-family:\'DM Mono\',monospace; font-size:0.68rem; '
    'color:#6B6860; margin-top:0.5rem;">'
    '🔒 Contracts are processed in memory — never stored, sold, or used to train AI models.'
    '</div>',
    unsafe_allow_html=True,
)

# Paste input — secondary path, expands automatically when sample is loaded
paste_expanded = bool(st.session_state.get("contract_text_input", "").strip())
with st.expander("Or paste contract text instead", expanded=paste_expanded):
    pasted = st.text_area(
        label="contract_text_input",
        label_visibility="collapsed",
        placeholder="Paste the full contract text here…",
        height=260,
        key="contract_text_input",
    )

# Zero-friction demo — populates the paste area with a predatory sample
st.button(
    "🧪 No contract handy? Try a sample predatory contract",
    on_click=_load_sample_contract,
    help="Loads a mock contractor agreement full of real-world traps so you can see the engine work.",
)

# Report structure preview — sets output expectation before first run
st.markdown(
    '<div style="font-family:\'DM Mono\',monospace; font-size:0.68rem; '
    'color:#6B6860; letter-spacing:0.04em; margin-top:1rem; line-height:1.8;">'
    'YOUR REPORT WILL INCLUDE &nbsp;→&nbsp; Contract Type &nbsp;·&nbsp; Risk Score '
    '&nbsp;·&nbsp; Top 3 Priority Issues &nbsp;·&nbsp; Red Flags with exact quotes '
    '&nbsp;·&nbsp; Compound Traps &nbsp;·&nbsp; Missing Protections '
    '&nbsp;·&nbsp; Negotiation Scripts &nbsp;·&nbsp; Sign / Negotiate / Avoid Verdict'
    '</div>',
    unsafe_allow_html=True,
)

contract_text: str | None = None
source_name: str = "unknown"
uploaded_text: str | None = None

if uploaded is not None:
    source_name = uploaded.name
    safe_uploaded_name = html.escape(uploaded.name)

    st.markdown(
        f'<div class="section-label" style="margin-top:0.5rem;">'
        f'File received&nbsp;&nbsp;<span style="color:#9A9690">{safe_uploaded_name}</span></div>',
        unsafe_allow_html=True,
    )

    with st.spinner("Extracting text…"):
        try:
            uploaded_text = extract_text_from_uploaded_file(uploaded)
        except RuntimeError as err:
            uploaded_text = None
            st.error(str(err))

    if uploaded_text and uploaded_text.strip():
        contract_text = uploaded_text.strip()
        source_name = uploaded.name
        st.success(
            f"Text extracted successfully from {uploaded.name} ({len(contract_text):,} characters)."
        )
    else:
        st.error(
            "We received the file, but could not extract readable contract text. "
            "Please paste the contract text manually instead."
        )

# File takes priority when extraction succeeds. If no usable file text exists,
# the paste area works exactly as before.
if not contract_text and pasted.strip():
    contract_text = pasted.strip()
    source_name = "pasted-text"

# ── Analyse button ────────────────────────────────────────────────────────────
st.markdown("<br>", unsafe_allow_html=True)

col_btn, col_spacer = st.columns([1, 5])
with col_btn:
    analyse_clicked = st.button(
        "Analyse Contract",
        disabled=(not contract_text or not (contract_text or "").strip()),
        use_container_width=False,
    )

# ── Run pipeline & render report ──────────────────────────────────────────────
if analyse_clicked and contract_text and contract_text.strip():
    st.markdown('<hr class="hairline">', unsafe_allow_html=True)

    with st.spinner("Running 4eyes analysis pipeline…"):
        try:
            report_output = run_analysis(contract_text)
            analysis_ok = True
        except Exception as exc:
            analysis_ok = False
            error_msg = str(exc)

    if analysis_ok:
        import datetime
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        safe_source_name = html.escape(source_name)

        # Meta bar
        st.markdown(f"""
        <div class="meta-bar">
            <div class="meta-item">Source&nbsp;&nbsp;<span class="meta-value">{safe_source_name}</span></div>
            <div class="meta-item">Length&nbsp;&nbsp;<span class="meta-value">{len(contract_text):,} chars</span></div>
            <div class="meta-item">Analysed&nbsp;&nbsp;<span class="meta-value">{ts}</span></div>
        </div>
        """, unsafe_allow_html=True)

        # Report
        st.markdown('<div class="section-label">Analysis Report</div>', unsafe_allow_html=True)

        with st.container():
            st.markdown('<div class="report-wrapper">', unsafe_allow_html=True)

            # Render as markdown if it looks like markdown, otherwise as code
            if any(c in report_output for c in ["##", "**", "- ", "\n#"]):
                st.markdown(report_output)
            else:
                # Plain text or JSON: render in a scrollable mono block
                st.code(report_output, language="markdown")

            st.markdown("</div>", unsafe_allow_html=True)

        # Download
        st.markdown("<br>", unsafe_allow_html=True)
        col_dl, _ = st.columns([1, 4])
        with col_dl:
            st.download_button(
                label="↓  Download report",
                data=report_output,
                file_name=f"4eyes_report_{Path(source_name).stem}.txt",
                mime="text/plain",
            )

    else:
        st.error(f"**Analysis failed.** {error_msg}")
        st.markdown(
            '<div class="section-label" style="margin-top:1rem;">Traceback</div>',
            unsafe_allow_html=True,
        )
        st.code(error_msg, language="python")

elif analyse_clicked:
    st.warning("No contract text to analyse. Upload a file or paste text above.")

# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown("""
<hr class="hairline" style="margin-top:4rem;">
<div style="font-family:'DM Mono',monospace; font-size:0.65rem; color:#3A3A45; text-align:center; padding-bottom:1rem;">
    4eyes.ai &nbsp;·&nbsp; Contract Intelligence &nbsp;·&nbsp; Beta
</div>
""", unsafe_allow_html=True)
